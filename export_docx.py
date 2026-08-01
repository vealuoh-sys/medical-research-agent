import csv
import json
import os
import sys

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def create_word_document(manuscript_text, references_text, output_filename="research_paper_draft.docx"):
    """
    Converts text manuscript and data tables into a styled MS Word (.docx) document.
    """
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, output_filename)
    
    doc = Document()
    
    # Configure 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Configure Normal style font (Calibri 11pt, dark gray)
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # Process text lines
    lines = manuscript_text.splitlines()
    in_references = False
    
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
            
        # Title
        if line_str.startswith("**Title:**") or line_str.startswith("Title:"):
            title_text = line_str.replace("**Title:**", "").replace("Title:", "").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_text)
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy blue
            doc.add_paragraph() # Spacer
            continue
            
        # Headings (e.g. **Abstract:**, **Introduction:**, **Methods:**, **Results:**, **Discussion and Limitations:**, **References:**)
        if (line_str.startswith("**") and line_str.endswith("**")) or (line_str.startswith("# ") or line_str.startswith("## ")):
            clean_heading = line_str.replace("**", "").replace("#", "").strip()
            if clean_heading.endswith(":"):
                clean_heading = clean_heading[:-1]
                
            if clean_heading.upper() == "REFERENCES":
                in_references = True
                
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(clean_heading)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            continue
            
        # Sub-headings / bold section markers
        if line_str.startswith("**") and ":" in line_str:
            parts = line_str.split(":", 1)
            sub_title = parts[0].replace("**", "").strip()
            sub_body = parts[1].strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            run_b = p.add_run(sub_title + ": ")
            run_b.font.bold = True
            run_b.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            p.add_run(sub_body)
            continue
            
        # List items
        if line_str.startswith("* ") or line_str.startswith("- "):
            item_text = line_str[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.add_run(item_text)
            continue
            
        # Skip plain text references if we are going to append programmatic references
        if in_references:
            continue
            
        # Regular Body Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(line_str)
        
    # --- TABLE 1: QUADAS-2 Risk of Bias Table (If risk_of_bias_table.csv exists) ---
    rob_csv_path = os.path.join(script_dir, "risk_of_bias_table.csv")
    if os.path.exists(rob_csv_path):
        doc.add_paragraph()
        h_table = doc.add_paragraph()
        h_table.paragraph_format.space_before = Pt(16)
        h_table.paragraph_format.space_after = Pt(6)
        r_t = h_table.add_run("Table 1: QUADAS-2 Methodological Risk of Bias Assessment")
        r_t.font.size = Pt(13)
        r_t.font.bold = True
        r_t.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        try:
            with open(rob_csv_path, "r", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                rob_rows = list(reader)
                
            if rob_rows:
                table = doc.add_table(rows=1, cols=6)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr_cells = table.rows[0].cells
                headers = ["PMID", "Patient Selection", "Index Test", "Reference Standard", "Flow & Timing", "Rationale"]
                for i, header_text in enumerate(headers):
                    hdr_cells[i].text = header_text
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    set_cell_background(hdr_cells[i], "1B365D") # Dark navy header
                    
                for r in rob_rows:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(r.get("pmid", ""))
                    row_cells[1].text = str(r.get("patient_selection_bias", ""))
                    row_cells[2].text = str(r.get("index_test_bias", ""))
                    row_cells[3].text = str(r.get("reference_standard_bias", ""))
                    row_cells[4].text = str(r.get("flow_timing_bias", ""))
                    row_cells[5].text = str(r.get("rationale", ""))
        except Exception as e:
            print(f"Notice: Could not build Word table for QUADAS-2: {e}")

    # --- APPEND PROGRAMMATIC REFERENCES SECTION ---
    if references_text:
        doc.add_paragraph()
        ref_h = doc.add_paragraph()
        ref_h.paragraph_format.space_before = Pt(18)
        ref_h.paragraph_format.space_after = Pt(8)
        r_ref = ref_h.add_run("References")
        r_ref.font.size = Pt(14)
        r_ref.font.bold = True
        r_ref.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
        for ref_line in references_text.splitlines():
            ref_line_str = ref_line.strip()
            if not ref_line_str or ref_line_str.startswith("## REFERENCES"):
                continue
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.space_after = Pt(4)
            p_ref.paragraph_format.line_spacing = 1.15
            p_ref.add_run(ref_line_str)

    try:
        doc.save(output_path)
        return True, output_path
    except Exception as e:
        print(f"Error saving MS Word document '{output_filename}': {e}")
        return False, None
